import os
import fitz
import docx


SUPPORTED_FORMATS = [".pdf", ".docx"]


def parse_pdf(file_path):
    extracted_text = ""

    try:
        pdf_document = fitz.open(file_path)
        total_pages = len(pdf_document)

        for page_number in range(total_pages):
            page = pdf_document[page_number]
            page_text = page.get_text()
            extracted_text += page_text

        pdf_document.close()

    except Exception as e:
        raise ValueError(f"Error reading PDF file: {str(e)}")

    cleaned_text = clean_text(extracted_text)
    return cleaned_text


def parse_docx(file_path):
    extracted_text = ""

    try:
        word_document = docx.Document(file_path)

        for paragraph in word_document.paragraphs:
            if paragraph.text.strip():
                extracted_text += paragraph.text + "\n"

        for table in word_document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        extracted_text += cell.text + "\n"

    except Exception as e:
        raise ValueError(f"Error reading DOCX file: {str(e)}")

    cleaned_text = clean_text(extracted_text)
    return cleaned_text


def clean_text(raw_text):
    if not raw_text:
        return ""

    lines = raw_text.split("\n")
    cleaned_lines = []

    for line in lines:
        stripped_line = line.strip()
        if stripped_line:
            cleaned_lines.append(stripped_line)

    final_text = "\n".join(cleaned_lines)
    return final_text


def get_file_extension(file_path):
    filename = os.path.basename(file_path)
    extension = os.path.splitext(filename)[1]
    extension = extension.lower()
    return extension


def parse_resume(file_path):
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    file_extension = get_file_extension(file_path)

    if file_extension not in SUPPORTED_FORMATS:
        raise ValueError(
            f"Unsupported file format: {file_extension}. "
            f"Please upload PDF or DOCX files only."
        )

    if file_extension == ".pdf":
        extracted_text = parse_pdf(file_path)
    elif file_extension == ".docx":
        extracted_text = parse_docx(file_path)

    if len(extracted_text) < 50:
        raise ValueError(
            "Resume text too short. "
            "File might be image-based or corrupted. "
            "Please upload a text-based PDF or DOCX."
        )

    return extracted_text